"""
Document generation service — Phase 5.

Generates three output files per approved site:
  1. letters_zip   — Letter A / B / C .docx files, zipped
  2. ukg_upload    — Payroll Metrics .xlsx for UKG import
  3. regional_excel — Approved-rates summary .xlsx for the regional manager
"""
from __future__ import annotations

import io
import re
import zipfile
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from app.models.cycle import ReviewCycle
    from app.models.employee import Employee


# ─────────────────────────────────────────────────────────────────────────────
#  Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _safe(text: str) -> str:
    """Strip characters invalid in filenames."""
    return re.sub(r'[\\/*?:"<>|]', "_", str(text or "unknown")).strip()


def _fmt_rate(val: float | None) -> str:
    try:
        return f"${float(val):.2f}"
    except (TypeError, ValueError):
        return "N/A"


def _fmt_date(d: date | None) -> str:
    if d is None:
        return ""
    return d.strftime("%d/%m/%Y")


# ─────────────────────────────────────────────────────────────────────────────
#  Letter builder (ported from POC generate_letters.py)
# ─────────────────────────────────────────────────────────────────────────────
def _build_letter_doc(
    letter_type: str,
    emp: "Employee",
    cycle: "ReviewCycle",
) -> bytes:
    """Build one pay letter as a .docx and return the raw bytes."""
    from docx import Document
    from docx.shared import Cm, Pt

    fy_label = cycle.fy_label or ""
    fy_parts = fy_label.replace("FY", "").split("-")
    fy_current = f"{fy_parts[0]}-{fy_parts[1]}" if len(fy_parts) == 2 else fy_label
    fy_prev = f"{int(fy_parts[0]) - 1}-{fy_parts[0]}" if fy_parts else ""

    letter_date = _fmt_date(cycle.letter_date)
    effective_text = f"first full pay period on or after {_fmt_date(cycle.effective_date)}"
    super_old = cycle.super_old or "11.5%"
    super_new = cycle.super_new or "12.0%"
    consultation_dl = (
        _fmt_date(cycle.consultation_deadline)
        if cycle.consultation_deadline
        else "COB 28 days from letter date"
    )
    signatory_name = cycle.signatory_name or "General Manager, Operations"
    signatory_title = cycle.signatory_title or ""
    signatory_company = cycle.signatory_company or "Carlisle Health"
    hr_email = cycle.hr_email or "peopleandculture@carlislehealth.com.au"

    first_name = emp.first_name or "[First Name]"
    current_rate = float(emp.current_rate) if emp.current_rate else 0.0
    proposed_rate = float(emp.proposed_rate) if emp.proposed_rate else 0.0
    award_level = emp.fy26_award or ""

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add(text: str, bold: bool = False, space_after: int = 6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        return p

    def blank():
        add("")

    # Header
    add(letter_date)
    add("Via Email Correspondence")
    blank()
    add(f"Dear {first_name}")
    blank()

    # Subject
    subjects = {
        "A": "Re: Increase in your Remuneration",
        "B": "Re: Increase in your Remuneration & Change to your Award Level",
        "C": "Re: Realignment of your Award level",
    }
    add(subjects[letter_type], bold=True)
    blank()

    # Body
    if letter_type == "A":
        add(
            f"In recognition of changes in Award rates and as per Carlisle Pay & Progression model "
            f"for {fy_current} and your continued commitment to the Carlisle Health Group, this letter "
            f"is to confirm the business will be increasing your remuneration."
        )
        add(
            f"Your Hourly Rate per your contract will be increased from {_fmt_rate(current_rate)} per hour "
            f"to {_fmt_rate(proposed_rate)} per hour (exclusive of superannuation). "
            f"Your level under the Award is: {award_level}"
        )
        add(
            f"This change will be reflected in your pay for the {effective_text}, "
            f"i.e. the first full pay period of the new financial year."
        )
    elif letter_type == "B":
        add(
            f"In recognition of changes in Award rates and as per Carlisle Pay & Progression model "
            f"for {fy_current} and your continued commitment to the Carlisle Health Group, this letter "
            f"is to confirm the business will be increasing your remuneration."
        )
        add(
            f"Your Hourly Rate per your contract will be increased from {_fmt_rate(current_rate)} per hour "
            f"to {_fmt_rate(proposed_rate)} per hour (exclusive of superannuation). This change will be "
            f"reflected in your pay for the {effective_text}, i.e. the first full pay period of the new "
            f"financial year."
        )
        add(
            f"We have simultaneously undertaken a review of all Award levels across the business to ensure "
            f"that the expectations of your role aligned to the most appropriate Award level. As a result of "
            f"this process Carlisle proposes to align your Award level to: {award_level}"
        )
        add(
            f"With this, we are asking for any feedback on the proposed Award level update to be sent to "
            f"Carlisle HR at {hr_email} by {consultation_dl}. If you do not have any feedback, "
            f"your Award level will be updated."
        )
    elif letter_type == "C":
        add(
            f"Carlisle has conducted a review of all Award levels across the business to ensure that the "
            f"expectations of your role aligned to the most appropriate Award level. As a result of this "
            f"process Carlisle proposes to align your Award level to: {award_level}"
        )
        add(
            f"As part of this proposed change, we are asking for any feedback on this to be sent to "
            f"Carlisle HR at {hr_email} by {consultation_dl} ('the consultation period'). Please take "
            f"note that the proposed alignment of your Award level does not impact your current payrate."
        )
        add(
            f"If you do not have any feedback, your Award level will be updated to the above."
        )

    blank()

    # Super paragraph
    add(
        f"The rate of superannuation will also increase in line with national requirements from "
        f"{super_old} for the {fy_prev} financial year to {super_new} for the {fy_current} "
        f"financial year. This applies to all employees."
    )
    blank()
    add(
        "All other terms and conditions governing your employment continue to apply as per "
        "your current employment contract."
    )
    blank()
    add(
        f"If you are unclear on anything contained in this correspondence, please contact HR on "
        f"{hr_email}."
    )
    blank()
    add(
        f"On behalf of the Carlisle Health Management Team and Board, we would like to thank you "
        f"for your ongoing contribution."
    )
    blank()
    add("Yours Sincerely")
    blank()
    blank()
    add(signatory_name, bold=True)
    if signatory_title:
        add(signatory_title)
    add(signatory_company)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  PDF letter builder  (draft — no Word dependency)
# ─────────────────────────────────────────────────────────────────────────────
def _build_letter_pdf(
    letter_type: str,
    emp: "Employee",
    cycle: "ReviewCycle",
    *,
    draft: bool = True,
) -> bytes:
    """Build one pay letter as a PDF and return the raw bytes.

    draft=True  → light-grey "DRAFT" watermark diagonally centred on the page.
    draft=False → clean final letter, no watermark.
    """
    from fpdf import FPDF

    fy_label   = cycle.fy_label or ""
    fy_parts   = fy_label.replace("FY", "").split("-")
    fy_current = f"{fy_parts[0]}-{fy_parts[1]}" if len(fy_parts) == 2 else fy_label
    fy_prev    = f"{int(fy_parts[0]) - 1}-{fy_parts[0]}" if fy_parts else ""

    letter_date    = _fmt_date(cycle.letter_date)
    effective_text = f"first full pay period on or after {_fmt_date(cycle.effective_date)}"
    super_old      = cycle.super_old or "11.5%"
    super_new      = cycle.super_new or "12.0%"
    consultation_dl = (
        _fmt_date(cycle.consultation_deadline)
        if hasattr(cycle, "consultation_deadline") and cycle.consultation_deadline
        else "COB 28 days from letter date"
    )
    signatory_name    = cycle.signatory_name or "General Manager, Operations"
    signatory_title   = cycle.signatory_title or ""
    signatory_company = cycle.signatory_company or "Carlisle Health"
    hr_email          = cycle.hr_email or "peopleandculture@carlislehealth.com.au"

    first_name    = emp.first_name or "[First Name]"
    current_rate  = float(emp.current_rate)  if emp.current_rate  else 0.0
    proposed_rate = float(emp.proposed_rate) if emp.proposed_rate else 0.0
    award_level   = emp.fy26_award or ""

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()
    pdf.set_margins(left=28, top=25, right=28)

    BRAND = (211, 46, 83)   # Carlisle brand red #d32e53
    DARK  = (30, 30, 30)
    GREY  = (100, 100, 100)

    def para(text: str, bold: bool = False, colour=DARK, size: int = 11, gap_after: float = 4):
        pdf.set_font("Helvetica", style="B" if bold else "", size=size)
        pdf.set_text_color(*colour)
        pdf.multi_cell(0, 5.5, text, align="L")
        if gap_after:
            pdf.ln(gap_after)

    def gap(h: float = 4):
        pdf.ln(h)

    # ── Logo top-right ────────────────────────────────────────────────────────
    # SVG lives in the sibling frontend package (4 directories up from here)
    _logo = Path(__file__).resolve().parent.parent.parent.parent / "frontend" / "public" / "carlisle-logo.svg"
    if _logo.exists():
        try:
            # Logo aspect ratio ≈ 4.14 : 1 → 44 mm wide ≈ 10.6 mm tall
            pdf.image(str(_logo), x=138, y=8, w=44)
        except Exception:
            pass  # silently degrade if SVG renderer can't handle the file

    # ── DRAFT watermark (preview only) ───────────────────────────────────────
    # Drawn BEFORE content so text renders on top of the watermark.
    # Pivot at exact A4 centre; offset by half string width for visual centring.
    if draft:
        _CX, _CY = 105.0, 148.5
        pdf.set_font("Helvetica", style="B", size=72)
        pdf.set_text_color(235, 235, 235)
        _dw = pdf.get_string_width("DRAFT")
        with pdf.rotation(angle=45, x=_CX, y=_CY):
            pdf.text(x=_CX - _dw / 2, y=_CY, text="DRAFT")

    # Position cursor below logo, leaving breathing room
    pdf.set_y(28)

    # Date + salutation
    para(letter_date, colour=GREY, size=10, gap_after=1)
    para("Via Email Correspondence", colour=GREY, size=10, gap_after=5)
    para(f"Dear {first_name}", size=11, gap_after=5)

    # Subject line
    subjects = {
        "A": "Re: Increase in your Remuneration",
        "B": "Re: Increase in your Remuneration & Change to your Award Level",
        "C": "Re: Realignment of your Award level",
    }
    para(subjects[letter_type], bold=True, colour=BRAND, size=11, gap_after=5)

    # Body paragraphs — same wording as the docx builder
    if letter_type == "A":
        para(
            f"In recognition of changes in Award rates and as per Carlisle Pay & Progression model "
            f"for {fy_current} and your continued commitment to the Carlisle Health Group, this letter "
            f"is to confirm the business will be increasing your remuneration."
        )
        para(
            f"Your Hourly Rate per your contract will be increased from {_fmt_rate(current_rate)} per hour "
            f"to {_fmt_rate(proposed_rate)} per hour (exclusive of superannuation). "
            f"Your level under the Award is: {award_level}"
        )
        para(
            f"This change will be reflected in your pay for the {effective_text}, "
            f"i.e. the first full pay period of the new financial year."
        )
    elif letter_type == "B":
        para(
            f"In recognition of changes in Award rates and as per Carlisle Pay & Progression model "
            f"for {fy_current} and your continued commitment to the Carlisle Health Group, this letter "
            f"is to confirm the business will be increasing your remuneration."
        )
        para(
            f"Your Hourly Rate per your contract will be increased from {_fmt_rate(current_rate)} per hour "
            f"to {_fmt_rate(proposed_rate)} per hour (exclusive of superannuation). This change will be "
            f"reflected in your pay for the {effective_text}, i.e. the first full pay period of the new "
            f"financial year."
        )
        para(
            f"We have simultaneously undertaken a review of all Award levels across the business to ensure "
            f"that the expectations of your role aligned to the most appropriate Award level. As a result of "
            f"this process Carlisle proposes to align your Award level to: {award_level}"
        )
        para(
            f"With this, we are asking for any feedback on the proposed Award level update to be sent to "
            f"Carlisle HR at {hr_email} by {consultation_dl}. If you do not have any feedback, "
            f"your Award level will be updated."
        )
    elif letter_type == "C":
        para(
            f"Carlisle has conducted a review of all Award levels across the business to ensure that the "
            f"expectations of your role aligned to the most appropriate Award level. As a result of this "
            f"process Carlisle proposes to align your Award level to: {award_level}"
        )
        para(
            f"As part of this proposed change, we are asking for any feedback on this to be sent to "
            f"Carlisle HR at {hr_email} by {consultation_dl} ('the consultation period'). Please take "
            f"note that the proposed alignment of your Award level does not impact your current payrate."
        )
        para("If you do not have any feedback, your Award level will be updated to the above.")

    gap(4)

    # Super paragraph (common to all)
    para(
        f"The rate of superannuation will also increase in line with national requirements from "
        f"{super_old} for the {fy_prev} financial year to {super_new} for the {fy_current} "
        f"financial year. This applies to all employees."
    )
    para(
        "All other terms and conditions governing your employment continue to apply as per "
        "your current employment contract."
    )
    para(
        f"If you are unclear on anything contained in this correspondence, please contact HR on "
        f"{hr_email}."
    )
    para(
        f"On behalf of the Carlisle Health Management Team and Board, we would like to thank you "
        f"for your ongoing contribution."
    )
    gap(6)
    para("Yours Sincerely")
    gap(14)
    para(signatory_name, bold=True, gap_after=1)
    if signatory_title:
        para(signatory_title, colour=GREY, size=10, gap_after=1)
    para(signatory_company, colour=GREY, size=10, gap_after=2)

    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────────────────────────
#  Public draft-letter helper
# ─────────────────────────────────────────────────────────────────────────────
def generate_draft_letter_pdf(emp: "Employee", cycle: "ReviewCycle") -> bytes:
    """Build a single-employee DRAFT pay-letter PDF (with watermark) for preview."""
    letter_type = (emp.letter_type or "").upper().strip()
    if letter_type not in ("A", "B", "C"):
        raise ValueError(f"Employee {emp.id} has no valid letter type assigned.")
    return _build_letter_pdf(letter_type, emp, cycle, draft=True)


def generate_final_letter_pdf(emp: "Employee", cycle: "ReviewCycle") -> bytes:
    """Build a single-employee final pay-letter PDF (no watermark) for distribution."""
    letter_type = (emp.letter_type or "").upper().strip()
    if letter_type not in ("A", "B", "C"):
        raise ValueError(f"Employee {emp.id} has no valid letter type assigned.")
    return _build_letter_pdf(letter_type, emp, cycle, draft=False)


# ─────────────────────────────────────────────────────────────────────────────
#  1. Letters ZIP
# ─────────────────────────────────────────────────────────────────────────────
def generate_letters_zip(
    employees: list["Employee"],
    cycle: "ReviewCycle",
    out_path: Path,
) -> int:
    """Generate final Letter A/B/C PDFs (no watermark), zip them, write to out_path.
    Returns number of letters generated."""

    MIN_INCREASE = 0.10  # skip trivial rate bumps

    buf = io.BytesIO()
    count = 0

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for emp in employees:
            if emp.is_departed:
                continue
            lt = (emp.letter_type or "").upper().strip()
            if lt not in ("A", "B", "C"):
                continue

            # Skip tiny increases for Letter A
            if lt == "A":
                cr = float(emp.current_rate or 0)
                pr = float(emp.proposed_rate or 0)
                if pr > 0 and cr > 0 and 0 < (pr - cr) < MIN_INCREASE:
                    continue

            pdf_bytes = _build_letter_pdf(lt, emp, cycle, draft=False)
            fname = f"{_safe(emp.last_name)}_{_safe(emp.first_name)}_Letter{lt}.pdf"
            zf.writestr(f"Letter {lt}/{fname}", pdf_bytes)
            count += 1

    out_path.write_bytes(buf.getvalue())
    return count


# ─────────────────────────────────────────────────────────────────────────────
#  2. UKG Payroll Metrics upload
# ─────────────────────────────────────────────────────────────────────────────
def generate_ukg_upload(
    employees: list["Employee"],
    cycle: "ReviewCycle",
    out_path: Path,
) -> int:
    """Generate the UKG Payroll Metrics import .xlsx. Returns row count."""
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    NAVY = "1F3864"; YELLOW = "FFF2CC"; WHITE = "FFFFFF"

    def fill(c): return PatternFill("solid", fgColor=c)
    def fnt(bold=False): return Font(bold=bold, size=10, name="Calibri")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Payroll Metrics"

    eff = _fmt_date(cycle.effective_date)

    headers = [
        ("Payroll Name",          True,  "Last, First — match UKG exactly"),
        ("Employee ID",           True,  "UKG Employee Number"),
        ("First Name",            True,  ""),
        ("Last Name",             True,  ""),
        ("New Hourly Rate",       True,  "Approved proposed rate"),
        ("Award Classification",  True,  "FY26 award level"),
        ("Effective Date",        True,  eff),
        ("Site",                  False, ""),
        ("Department",            False, ""),
        ("Letter Type",           False, "For reference"),
    ]

    col_widths = [28, 15, 15, 15, 16, 24, 14, 20, 20, 12]

    # Row 1: title banner
    ws.row_dimensions[1].height = 20
    title_text = (
        f"UKG Payroll Metrics Upload — {cycle.fy_label}  |  "
        f"Effective {eff}  |  Yellow = mandatory"
    )
    tc = ws.cell(1, 1, title_text)
    tc.font = Font(bold=True, size=11, color=WHITE, name="Calibri")
    tc.fill = fill(NAVY)
    ws.merge_cells(f"A1:{chr(64 + len(headers))}1")

    # Row 2: column headers
    ws.row_dimensions[2].height = 36
    for i, (h, mandatory, note) in enumerate(headers, 1):
        c = ws.cell(2, i, h + (" *" if mandatory else ""))
        c.font = Font(bold=True, size=10, name="Calibri")
        c.fill = fill(YELLOW if mandatory else "F2F2F2")
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.column_dimensions[chr(64 + i)].width = col_widths[i - 1]

    # Data rows
    row_num = 3
    for emp in employees:
        if emp.is_departed:
            continue
        if not emp.proposed_rate:
            continue

        payroll_name = f"{emp.last_name}, {emp.first_name}".strip(", ")
        rate = round(float(emp.proposed_rate), 4)
        bg = "F2F2F2" if row_num % 2 == 0 else "FFFFFF"

        row_data = [
            payroll_name,
            emp.emp_num,
            emp.first_name,
            emp.last_name,
            rate,
            emp.fy26_award or "",
            eff,
            emp.site,
            emp.department or "",
            emp.letter_type or "",
        ]
        for i, val in enumerate(row_data, 1):
            c2 = ws.cell(row_num, i, val)
            c2.fill = fill(bg)
            c2.font = fnt()
            if i == 5:
                c2.number_format = '"$"#,##0.0000'
        row_num += 1

    wb.save(str(out_path))
    return row_num - 3


# ─────────────────────────────────────────────────────────────────────────────
#  3. Regional approved-rates summary Excel
# ─────────────────────────────────────────────────────────────────────────────
def generate_regional_excel(
    employees: list["Employee"],
    cycle: "ReviewCycle",
    site: str,
    out_path: Path,
) -> int:
    """Generate a read-only approved-rates summary Excel for the regional manager."""
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    NAVY = "1F3864"; WHITE = "FFFFFF"; GREY = "F2F2F2"; GREEN = "E2EFDA"
    AMBER = "FFF2CC"; BLUE = "D6E4F0"

    def fill(c): return PatternFill("solid", fgColor=c)
    def fnt(bold=False, color="000000", size=10):
        return Font(bold=bold, color=color, size=size, name="Calibri")
    def thin():
        s = Side(style="thin", color="CCCCCC")
        return Border(left=s, right=s, top=s, bottom=s)
    def centre(): return Alignment(horizontal="center", vertical="center")
    def left(): return Alignment(horizontal="left", vertical="center")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Approved Rates"
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A3"

    columns = [
        ("Emp #",           10),
        ("First Name",      16),
        ("Last Name",       16),
        ("Site",            18),
        ("Department",      20),
        ("Category",        14),
        ("Hrs/Wk",           8),
        ("FY25 Award",      20),
        ("Current Rate",    13),
        ("FY26 Award",      20),
        ("PP Level",        32),
        ("Proposed Rate",   13),
        ("$ Change",        11),
        ("% Change",        10),
        ("Change Type",     16),
        ("Letter",           9),
        ("Notes",           28),
    ]

    # Row 1: title
    ws.row_dimensions[1].height = 26
    tc = ws.cell(1, 1, f"Approved Pay Rates — {site} — {cycle.fy_label}  |  Effective {_fmt_date(cycle.effective_date)}")
    tc.font = Font(bold=True, size=12, color=WHITE, name="Calibri")
    tc.fill = fill(NAVY)
    tc.alignment = left()
    ws.merge_cells(f"A1:{get_column_letter(len(columns))}1")

    # Row 2: headers
    ws.row_dimensions[2].height = 36
    for i, (name, width) in enumerate(columns, 1):
        c = ws.cell(2, i, name)
        c.font = fnt(bold=True, color=WHITE, size=9)
        c.fill = fill(NAVY)
        c.alignment = centre()
        c.border = thin()
        ws.column_dimensions[get_column_letter(i)].width = width

    # Data
    active = [e for e in employees if not e.is_departed]
    for row_num, emp in enumerate(active, start=3):
        ws.row_dimensions[row_num].height = 17
        bg = GREY if row_num % 2 == 1 else WHITE

        cr = float(emp.current_rate or 0)
        pr = float(emp.proposed_rate or 0)
        delta = pr - cr if pr and cr else None
        pct = (delta / cr * 100) if delta is not None and cr else None

        lt = (emp.letter_type or "").upper()
        lt_bg = AMBER if lt == "B" else (BLUE if lt == "C" else bg)

        row_data = [
            emp.emp_num,
            emp.first_name,
            emp.last_name,
            emp.site,
            emp.department or "",
            emp.category or "",
            emp.hours_per_week,
            emp.fy25_award or "",
            cr if cr else None,
            emp.fy26_award or "",
            emp.pp_level or "",
            pr if pr else None,
            delta,
            pct,
            emp.change_type or "",
            emp.letter_type or "",
            emp.notes or "",
        ]

        for i, val in enumerate(row_data, 1):
            use_bg = lt_bg if i == 16 else bg
            c = ws.cell(row_num, i, val)
            c.fill = fill(use_bg)
            c.font = fnt(size=10)
            c.border = thin()
            c.alignment = left()
            col_name = columns[i - 1][0]
            if col_name in ("Current Rate", "Proposed Rate", "$ Change"):
                c.number_format = '"$"#,##0.00'
            elif col_name == "% Change" and val is not None:
                c.value = round(val / 100, 4)
                c.number_format = "0.00%"
            elif col_name in ("Hrs/Wk",):
                c.number_format = "#,##0.##"

    # Summary at bottom
    last_data = 2 + len(active)
    ws.row_dimensions[last_data + 2].height = 20
    s1 = ws.cell(last_data + 2, 1, f"Total employees: {len(active)}")
    s1.font = fnt(bold=True)
    if active:
        total_curr = sum(float(e.current_rate or 0) * float(e.hours_per_week or 0) * 52 for e in active)
        total_prop = sum(float(e.proposed_rate or 0) * float(e.hours_per_week or 0) * 52 for e in active)
        ws.cell(last_data + 2, 9, total_curr).number_format = '"$"#,##0'
        ws.cell(last_data + 2, 12, total_prop).number_format = '"$"#,##0'
        ws.cell(last_data + 2, 9).font = fnt(bold=True)
        ws.cell(last_data + 2, 12).font = fnt(bold=True)

    wb.save(str(out_path))
    return len(active)
